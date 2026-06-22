"""
Parse uploaded manifest files (.xlsx, .docx, .txt) and extract field values.
Matching is case-insensitive on the field label.
"""
import io
from typing import Dict, List, Optional


def parse_uploaded_file(uploaded_file, questions: List[Dict]) -> Optional[Dict[str, str]]:
    """
    Detect file type and delegate to the appropriate parser.
    Returns {field_key: value} or None on failure.
    """
    name = uploaded_file.name.lower()
    try:
        if name.endswith(".xlsx"):
            return _parse_excel(uploaded_file, questions)
        if name.endswith(".docx"):
            return _parse_docx(uploaded_file, questions)
        if name.endswith(".txt"):
            return _parse_txt(uploaded_file, questions)
        return None
    except Exception as exc:
        print(f"[file_parser] Error parsing '{uploaded_file.name}': {exc}")
        return None


def _label_map(questions: List[Dict]) -> Dict[str, str]:
    """Build a lower-case label → field key lookup."""
    return {q["label"].lower(): q["key"] for q in questions}


def _parse_excel(uploaded_file, questions: List[Dict]) -> Dict[str, str]:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(uploaded_file.read()))
    ws = wb.active
    lmap = _label_map(questions)
    result: Dict[str, str] = {}
    for row in ws.iter_rows(values_only=True):
        if row[0] is not None and len(row) >= 2:
            label = str(row[0]).strip().lower()
            value = str(row[1]).strip() if row[1] is not None else ""
            key = lmap.get(label)
            if key:
                result[key] = value
    return result


def _parse_docx(uploaded_file, questions: List[Dict]) -> Dict[str, str]:
    from docx import Document
    doc = Document(io.BytesIO(uploaded_file.read()))
    lmap = _label_map(questions)
    result: Dict[str, str] = {}
    for para in doc.paragraphs:
        if ":" in para.text:
            label, _, value = para.text.partition(":")
            key = lmap.get(label.strip().lower())
            if key:
                result[key] = value.strip()
    # Also check table cells (template uses a table)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                key = lmap.get(cells[0].lower())
                if key and cells[1]:
                    result[key] = cells[1]
    return result


def _parse_txt(uploaded_file, questions: List[Dict]) -> Dict[str, str]:
    content = uploaded_file.read().decode("utf-8", errors="replace")
    lmap = _label_map(questions)
    result: Dict[str, str] = {}
    for line in content.splitlines():
        if ":" in line:
            label, _, value = line.partition(":")
            key = lmap.get(label.strip().lower())
            if key:
                result[key] = value.strip()
    return result
