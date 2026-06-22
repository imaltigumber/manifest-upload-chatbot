"""
Generate manifest files (Excel, Word, Text) — both filled and blank templates.
"""
import io
from datetime import datetime
from typing import Dict, List


# ── Helpers ───────────────────────────────────────────────────────────────────

def _ts() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M")


def _safe_name(intent_name: str) -> str:
    return intent_name.replace(" ", "_")


# ── Filled manifests ──────────────────────────────────────────────────────────

def generate_excel(answers: Dict[str, str], questions: List[Dict],
                   intent_name: str, version: int) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Manifest"

    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin"),
    )

    # Title row
    ws.merge_cells("A1:B1")
    ws["A1"] = f"Client Onboarding Manifest — {intent_name}  (v{version})"
    ws["A1"].font      = Font(bold=True, color="FFFFFF", size=13)
    ws["A1"].fill      = PatternFill("solid", fgColor="1E3A8A")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    # Generated timestamp
    ws["A2"] = "Generated:"
    ws["B2"] = _ts()
    ws["A2"].font = Font(italic=True, color="555555")

    # Column headers
    ws["A3"] = "Field"
    ws["B3"] = "Value"
    for cell in (ws["A3"], ws["B3"]):
        cell.font   = Font(bold=True)
        cell.fill   = PatternFill("solid", fgColor="D9E1F2")
        cell.border = thin

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 42

    # Data rows
    for r, q in enumerate(questions, start=4):
        a = ws.cell(row=r, column=1, value=q["label"])
        b = ws.cell(row=r, column=2, value=answers.get(q["key"], ""))
        a.border = thin
        b.border = thin
        if r % 2 == 0:
            a.fill = PatternFill("solid", fgColor="F8FAFC")
            b.fill = PatternFill("solid", fgColor="F8FAFC")

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def generate_docx(answers: Dict[str, str], questions: List[Dict],
                  intent_name: str, version: int) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn
    import copy

    doc = Document()
    doc.add_heading("Client Onboarding Manifest", 0)
    doc.add_heading(f"{intent_name}  —  Version {version}", 1)

    p = doc.add_paragraph()
    run = p.add_run(f"Generated: {_ts()}")
    run.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for q in questions:
        row = table.add_row().cells
        row[0].text = q["label"]
        row[1].text = str(answers.get(q["key"], ""))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_txt(answers: Dict[str, str], questions: List[Dict],
                 intent_name: str, version: int) -> str:
    lines = [
        "=" * 52,
        "  CLIENT ONBOARDING MANIFEST",
        f"  {intent_name}  |  Version {version}",
        f"  Generated: {_ts()}",
        "=" * 52,
        "",
    ]
    for q in questions:
        lines.append(f"{q['label']}: {answers.get(q['key'], '')}")
    lines += ["", "=" * 52]
    return "\n".join(lines)


# ── Blank templates ───────────────────────────────────────────────────────────

def generate_template_excel(questions: List[Dict], intent_name: str) -> bytes:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Manifest Template"

    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"),  bottom=Side(style="thin"),
    )

    ws.merge_cells("A1:B1")
    ws["A1"] = f"Manifest Template — {intent_name}"
    ws["A1"].font = Font(bold=True, color="FFFFFF", size=12)
    ws["A1"].fill = PatternFill("solid", fgColor="1E3A8A")

    ws["A2"] = "Field"
    ws["B2"] = "Value  (fill in and upload)"
    for cell in (ws["A2"], ws["B2"]):
        cell.font   = Font(bold=True)
        cell.fill   = PatternFill("solid", fgColor="D9E1F2")
        cell.border = thin

    ws.column_dimensions["A"].width = 28
    ws.column_dimensions["B"].width = 42

    for r, q in enumerate(questions, start=3):
        a = ws.cell(row=r, column=1, value=q["label"])
        b = ws.cell(row=r, column=2, value="")
        a.border = thin
        b.border = thin

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def generate_template_docx(questions: List[Dict], intent_name: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(f"Manifest Template — {intent_name}", 0)
    doc.add_paragraph(
        "Fill in the value after each colon, save the file, then upload it back."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Value  (fill in)"
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True

    for q in questions:
        row = table.add_row().cells
        row[0].text = q["label"]
        row[1].text = ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generate_template_txt(questions: List[Dict], intent_name: str) -> str:
    lines = [
        "=" * 52,
        f"  MANIFEST TEMPLATE — {intent_name}",
        "  Fill in the values after each colon.",
        "  Save the file and upload it back.",
        "=" * 52,
        "",
    ]
    for q in questions:
        lines.append(f"{q['label']}: ")
    lines += ["", "=" * 52]
    return "\n".join(lines)
